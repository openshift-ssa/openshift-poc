#!/usr/bin/env python3
"""Generate the POC Checklist as a Word document.

Run during CI to keep the .docx in sync with the markdown checklist.
Output: docs/assets/downloads/poc-checklist.docx

Styling follows Red Hat brand standards:
- Fonts: Red Hat Display (headings), Red Hat Text (body)
- Colors: Red Hat red #EE0000 for accents, black/white/gray for structure
- Layout: Generous margins, full-width tables, sentence case
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

RH_RED = RGBColor(0xEE, 0x00, 0x00)
RH_BLACK = RGBColor(0x15, 0x15, 0x15)
RH_GRAY_70 = RGBColor(0x38, 0x38, 0x38)
RH_GRAY_20 = RGBColor(0xE0, 0xE0, 0xE0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT_DISPLAY = 'Red Hat Display'
FONT_TEXT = 'Red Hat Text'

PAGE_WIDTH = Cm(21.0)
LEFT_MARGIN = Cm(1.5)
RIGHT_MARGIN = Cm(1.5)
TABLE_WIDTH = PAGE_WIDTH - LEFT_MARGIN - RIGHT_MARGIN  # 18cm usable

doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = PAGE_WIDTH
    section.page_height = Cm(29.7)
    section.left_margin = LEFT_MARGIN
    section.right_margin = RIGHT_MARGIN
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# Style: Normal (body text)
style = doc.styles['Normal']
style.font.name = FONT_TEXT
style.font.size = Pt(10)
style.font.color.rgb = RH_BLACK
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.3

# Style: Heading 1 (phase headings)
h1 = doc.styles['Heading 1']
h1.font.name = FONT_DISPLAY
h1.font.size = Pt(18)
h1.font.bold = True
h1.font.color.rgb = RH_RED
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(8)

# Style: Heading 2 (section headings)
h2 = doc.styles['Heading 2']
h2.font.name = FONT_DISPLAY
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RH_GRAY_70
h2.paragraph_format.space_before = Pt(16)
h2.paragraph_format.space_after = Pt(6)


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, font_color=WHITE, bg_hex='EE0000'):
    """Style a header row with brand colors."""
    for cell in row.cells:
        set_cell_shading(cell, bg_hex)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = font_color
                run.font.name = FONT_DISPLAY
                run.font.size = Pt(9)


def set_col_widths(table, widths):
    """Set explicit column widths, grid, and force fixed table layout."""
    from docx.oxml.ns import qn as _qn
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    # Convert Cm/Emu widths to twips (1 twip = 635 EMU)
    widths_twips = [int(w / 635) for w in widths]
    total_twips = sum(widths_twips)

    # Set table width to exact value
    tbl_w = tbl_pr.find(_qn('w:tblW'))
    if tbl_w is None:
        tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_twips}" w:type="dxa"/>')
        tbl_pr.append(tbl_w)
    else:
        tbl_w.set(_qn('w:w'), str(total_twips))
        tbl_w.set(_qn('w:type'), 'dxa')

    # Set fixed layout
    tbl_layout = tbl_pr.find(_qn('w:tblLayout'))
    if tbl_layout is None:
        tbl_layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
        tbl_pr.append(tbl_layout)

    # Update tblGrid to match column widths
    tbl_grid = tbl.find(_qn('w:tblGrid'))
    if tbl_grid is not None:
        tbl.remove(tbl_grid)
    tbl_grid = parse_xml(f'<w:tblGrid {nsdecls("w")}/>'.encode())
    for w in widths_twips:
        grid_col = parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{w}"/>')
        tbl_grid.append(grid_col)
    tbl.insert(tbl.index(tbl_pr) + 1, tbl_grid)

    # Set each cell width explicitly
    for row in table.rows:
        for i, width in enumerate(widths):
            cell = row.cells[i]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(_qn('w:tcW'))
            w_twips = widths_twips[i]
            if tc_w is None:
                tc_w = parse_xml(
                    f'<w:tcW {nsdecls("w")} w:w="{w_twips}" w:type="dxa"/>'
                )
                tc_pr.append(tc_w)
            else:
                tc_w.set(_qn('w:w'), str(w_twips))
                tc_w.set(_qn('w:type'), 'dxa')


def style_body_cells(table, start_row=1):
    """Apply body font styling to non-header rows."""
    for row in table.rows[start_row:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = FONT_TEXT
                    run.font.size = Pt(9)
                    run.font.color.rgb = RH_BLACK


def add_phase_heading(text):
    doc.add_heading(text, level=1)


def add_section_heading(text):
    doc.add_heading(text, level=2)


def add_checklist_table(items):
    """Standard checklist: Item, Status, Notes — full page width."""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    hdr[0].text = 'Item'
    hdr[1].text = 'Status'
    hdr[2].text = 'Notes'
    style_header_row(table.rows[0])

    for item in items:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = '☐'
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].text = ''

    col_widths = [Cm(9.0), Cm(1.8), Cm(7.2)]
    set_col_widths(table, col_widths)
    style_body_cells(table)
    doc.add_paragraph()


def add_validation_table(rows_data):
    """Exit criteria table: Area, How We Measure Success, Status, Notes."""
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    hdr[0].text = 'Area'
    hdr[1].text = 'How we measure success'
    hdr[2].text = 'Status'
    hdr[3].text = 'Notes'
    style_header_row(table.rows[0])

    for area, criteria in rows_data:
        row = table.add_row().cells
        row[0].text = area
        row[1].text = criteria
        row[2].text = ''
        row[3].text = ''

    col_widths = [Cm(2.5), Cm(9.0), Cm(1.5), Cm(5.0)]
    set_col_widths(table, col_widths)
    style_body_cells(table)
    doc.add_paragraph()


def add_results_table(rows_data):
    """Results summary: Category, What Was Tested, Expected, Actual, Result."""
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'What was tested'
    hdr[2].text = 'Expected outcome'
    hdr[3].text = 'Actual outcome'
    hdr[4].text = 'Result'
    style_header_row(table.rows[0])

    for row_data in rows_data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    col_widths = [Cm(2.5), Cm(3.5), Cm(4.5), Cm(4.5), Cm(3.0)]
    set_col_widths(table, col_widths)
    style_body_cells(table)
    doc.add_paragraph()


def add_signoff_table():
    """Formal sign-off table with role, name, signature, date."""
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    hdr[0].text = 'Role'
    hdr[1].text = 'Name'
    hdr[2].text = 'Signature'
    hdr[3].text = 'Date'
    style_header_row(table.rows[0])

    roles = [
        'Customer Technical Lead',
        'Customer Infrastructure Lead',
        'Red Hat / Partner Engineer',
        'Project Sponsor',
    ]
    for role in roles:
        row = table.add_row().cells
        row[0].text = role

    col_widths = [Cm(5.0), Cm(4.5), Cm(5.5), Cm(3.0)]
    set_col_widths(table, col_widths)
    style_body_cells(table)
    doc.add_paragraph()


# --- Document content ---

title = doc.add_heading('OpenShift POC checklist', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = FONT_DISPLAY
    run.font.size = Pt(24)
    run.font.color.rgb = RH_RED

doc.add_paragraph(
    'This checklist provides a complete end-to-end baseline for validating OpenShift '
    'in your environment. Not every item will apply to every environment \u2014 skip sections '
    'that are out of scope for your specific goals.'
)
doc.add_paragraph()

# Phase 1
add_phase_heading('Phase 1: Discovery and scoping')
doc.add_paragraph(
    'Complete this phase before any technical work begins. '
    'Align on goals, boundaries, and what a successful outcome looks like.'
)

add_section_heading('Goals and drivers')
add_checklist_table([
    'Document specific POC objectives with measurable outcomes',
    'Identify which applications or workloads will be evaluated',
    'Capture high-availability and recovery expectations',
    'Note any security, regulatory, or compliance constraints',
    'Understand target production timeline and anticipated scale',
])

add_section_heading('Boundaries and assumptions')
add_checklist_table([
    'Agree on what is included in the POC and document it',
    'Explicitly list what is excluded (to prevent scope creep)',
    'Document assumptions and responsibilities (customer vs. Red Hat)',
    'Identify all participating teams (infra, network, security, dev)',
    'Set a timeline with key milestone dates',
])

add_section_heading('Exit criteria')
doc.add_paragraph('Agree on pass/fail criteria before installation begins:')
add_validation_table([
    ('Installation', 'Cluster deployed, all nodes healthy'),
    ('Networking', 'Application traffic flows end-to-end'),
    ('Storage', 'Persistent volumes provision and bind on demand'),
    ('Security', 'Users authenticate and RBAC enforces permissions'),
    ('Applications', 'Sample and customer workloads run correctly'),
    ('Resilience', 'Cluster recovers from simulated failures'),
    ('Operations', 'Monitoring, logging, and backups function'),
    ('Migration', 'VMs migrated to OpenShift (if applicable)'),
])

# Phase 2
add_phase_heading('Phase 2: Prerequisites')
doc.add_paragraph('Complete these before scheduling the installation.')

add_section_heading('Infrastructure and compute')
add_checklist_table([
    'Red Hat account created and subscriptions allocated',
    'Compute nodes provisioned and meet minimum specs',
    'Management access to nodes confirmed (BMC, vCenter, cloud API)',
    'Firmware/BIOS settings validated (UEFI, virtualization extensions)',
    'Node details collected (MAC addresses, BMC IPs, disk hints, NIC names)',
])

add_section_heading('Networking')
add_checklist_table([
    'Network subnets allocated for cluster traffic',
    'Firewall rules opened (API, ingress, registry access)',
    'Outbound connectivity to Red Hat registries and services verified',
    'HTTP proxy configured and CA bundle prepared (if proxied environment)',
    'NTP accessible from all nodes',
    'API VIP and Ingress VIP assigned (bare metal) or user-managed load balancer configured',
])

add_section_heading('DNS')
add_checklist_table([
    'api.<cluster>.<domain> DNS record resolves correctly',
    'api-int.<cluster>.<domain> DNS record resolves correctly',
    '*.apps.<cluster>.<domain> wildcard DNS resolves correctly',
    'Reverse DNS entries (PTR) configured for node IPs',
])

add_section_heading('Storage')
add_checklist_table([
    'Storage array make, model, and firmware version documented',
    'Connection protocol identified (FC, iSCSI, NVMe/TCP, NVMe/FC, NFS)',
    'Array type documented (ALUA, active/active symmetric, active/passive)',
    'Number of fabric paths per node confirmed (minimum two for redundancy)',
    'FC: HBA WWPNs collected for all nodes',
    'FC: Zoning configured on both fabric switches',
    'FC: LUN masking / host groups configured on the array',
    'iSCSI: Target portal IPs and IQN documented',
    'iSCSI: CHAP credentials obtained (if required)',
    'NVMe-oF: Discovery controller IPs or target NQNs documented',
    'NVMe-oF: Array firmware confirmed to support NVMe-oF with RHCOS',
    'Dedicated storage VLAN ID and subnet allocated',
    'Jumbo frames (MTU 9000) confirmed end-to-end on the storage network',
    'Vendor-recommended multipath.conf device block obtained',
    'Multi-vendor: all vendor device blocks collected for merged multipath.conf',
    'Storage array accessible from all cluster node networks',
    'Required firewall ports open between nodes and the storage array',
    'Storage credentials or certificates available for CSI driver configuration',
    'RWX support confirmed (required for Virtualization live migration + registry)',
    'CSI driver version confirmed compatible with target OCP version',
])

add_section_heading('Installation host')
add_checklist_table([
    'oc CLI installed on installation host',
    'Pull secret downloaded from Red Hat console',
    'SSH key pair generated',
    'Disconnected: registry set up (oc-mirror or pull-through cache)',
    'Disconnected: install-config imageContentSources and OLM catalogs pointed at the registry',
])

add_section_heading('VM migration prerequisites')
add_checklist_table([
    'VDDK access requested from Broadcom (support ticket required)',
    'VDDK archive downloaded and version matched to vSphere version',
])

# Phase 3
add_phase_heading('Phase 3: Installation')
add_checklist_table([
    'Installation method selected',
    'Cluster installation completed successfully',
    'All nodes showing Ready status',
    'All ClusterOperators Available=True, Degraded=False',
    'Cluster version matches target release',
    'Web console accessible',
    'kubeadmin credentials stored securely',
])

# Phase 4
add_phase_heading('Phase 4: Configure the cluster (required)')
doc.add_paragraph('These must be completed before deploying workloads.')
add_checklist_table([
    'NMState operator installed',
    'Network configuration applied (bonds, VLANs, storage network NNCPs)',
    'Jumbo frames (MTU 9000) verified end-to-end on the storage network',
    'Storage driver installed and StorageClasses created',
    'Default StorageClass set',
    'Multipath configured and verified (iSCSI/FC only \u2014 MachineConfig applied)',
    'RWO PVC created, bound, and data write/read verified',
    'RWX PVC created and bound successfully (if applicable)',
    'Internal image registry configured with persistent storage',
])

# Phase 5
add_phase_heading('Phase 5: Configure the cluster (additional)')
doc.add_paragraph('Install based on your POC goals. Each subsection is independent.')

add_section_heading('Networking')
add_checklist_table([
    'OVS bridge trunk configured for VM secondary networks (if applicable)',
    'ClusterUserDefinedNetwork (CUDN) created for VM IPAM (if applicable)',
    'Ingress/Route exposes an application externally',
    'DNS resolution works from within pods (internal and external)',
])

add_section_heading('Workload availability')
add_checklist_table([
    'Node Health Check operator installed',
    'NodeHealthCheck CRs created (workers and control plane)',
    'Self Node Remediation operator installed',
    'Kube Descheduler operator installed and configured',
])

add_section_heading('Virtualization')
add_checklist_table([
    'OpenShift Virtualization operator installed',
    'HyperConverged CR created',
    'Virtualization StorageClass annotated as default virt class',
    'Live migration network configured (if applicable)',
])

add_section_heading('Migration')
add_checklist_table([
    'Migration Toolkit for Virtualization (MTV) operator installed',
    'VDDK image built and pushed to registry',
    'Source virtualization provider added and healthy',
    'Network and storage mappings configured',
])

add_section_heading('Backup and restore')
add_checklist_table([
    'OADP (OpenShift API for Data Protection) operator installed',
    'BackupStorageLocation CR created and showing Available',
    'DataProtectionApplication CR created',
])

add_section_heading('Observability')
add_checklist_table([
    'Logging operators installed (Loki Operator, OpenShift Logging)',
    'LokiStack deployed with object storage backend',
    'Log forwarding configured (ClusterLogForwarder)',
    'Network observability operator enabled',
    'Multi-cluster observability configured (if multi-cluster)',
])

add_section_heading('GitOps')
add_checklist_table([
    'OpenShift GitOps operator installed',
    'ArgoCD instance accessible',
    'Sample application deployed via GitOps',
])

add_section_heading('Service Mesh (out of baseline)')
doc.add_paragraph(
    'Skip unless mesh (mTLS, traffic splitting) is explicitly in POC scope.'
)
add_checklist_table([
    'OpenShift Service Mesh 3.x installed (Istio ambient)',
    'Bookinfo (or equivalent) deployed and mTLS verified',
])

add_section_heading('Security and access')
add_checklist_table([
    'Identity provider configured (LDAP, OIDC, etc.) \u2014 before demo day; keep kubeadmin until an IdP user has cluster-admin',
    'RBAC groups mapped correctly (members get expected roles)',
    'kubeadmin secret removed after IdP verification (if desired)',
    'External Secrets Operator installed (if integrating Vault, AWS Secrets Manager, etc.)',
    'Non-production banner applied to the console',
])

add_section_heading('Additional tools')
add_checklist_table([
    'Web Terminal operator installed (embedded CLI in the web console)',
    'Operators from Artifactory configured (if private registry for catalog)',
])

# Phase 6
add_phase_heading('Phase 6: VM migration')
doc.add_paragraph(
    'Validate that virtual machines can be migrated from an existing '
    'virtualization platform to OpenShift Virtualization.'
)
add_checklist_table([
    'Source virtualization environment accessible from OpenShift',
    'Migration provider connection healthy',
    'Target storage class selected',
    'Target network mapping configured',
    'VMs selected for migration',
    'Cold migration executed \u2014 VM boots on OpenShift',
    'Networking functional (IP, DNS, connectivity)',
    'Storage attached and data intact',
    'Applications inside the VM running correctly',
    'Warm migration tested \u2014 cutover with minimal downtime',
    'VM accessible via console and SSH post-migration',
])

# Phase 7
add_phase_heading('Phase 7: Workloads')
doc.add_paragraph('Deploy workloads to validate platform capabilities.')

add_section_heading('Container workloads')
add_checklist_table([
    'Basic container deployed and accessible via Route',
    'Build from source \u2014 image built and app deploys',
    'Stateful application \u2014 data persists across pod restarts',
    'Multi-tier application \u2014 frontend and backend communicating',
    'Event streaming workload deployed, e.g. Kafka (if applicable)',
    'Customer application deployed (if provided)',
])

add_section_heading('Virtual machine workloads')
add_checklist_table([
    'VM deployed from template \u2014 boots, SSH, storage functional',
    'Live migration tested (move VM between nodes without downtime)',
    'Snapshot and restore tested',
    'OVA virtual appliance deployed (if applicable)',
])

# Phase 8
add_phase_heading('Phase 8: Operational validation')
doc.add_paragraph('Demonstrate Day 2 operations and resilience.')

add_section_heading('Failover and resilience')
add_checklist_table([
    'Node failure simulated',
    'VM restarted on healthy node within target time',
    'Application recovered without manual intervention',
    'Container pod rescheduled to healthy node after failure',
    'Service remained available during failover',
])

add_section_heading('Backup and restore')
add_checklist_table([
    'Application or VM backup completed successfully',
    'Restore to same or different namespace validated',
    'Data integrity confirmed after restore',
])

add_section_heading('Cluster lifecycle')
add_checklist_table([
    'SSH key rotation validated',
    'Node-level configuration change applied via MachineConfig',
    'Node drain and maintenance \u2014 cordon, drain, uncordon',
    'Cluster upgrade tested (minor version or z-stream)',
    'Workloads remained available during upgrade',
    'Worker node added \u2014 new node joins cluster successfully',
])

add_section_heading('Monitoring and troubleshooting')
add_checklist_table([
    'Monitoring dashboards accessible',
    'Alerts fire correctly (trigger test alert, verify delivery)',
    'MTU verified end-to-end (no silent packet drops)',
    'must-gather diagnostic bundle collected and reviewed',
    'Log collection validated (node, pod, operator logs)',
    'Common failure modes understood by the customer team',
])

# Phase 9
add_phase_heading('Phase 9: Results summary')
doc.add_paragraph(
    'Complete this table during the POC to prepare for the closeout meeting readout.'
)
add_results_table([
    ('Installation', 'Cluster deploy', 'All nodes healthy', '', ''),
    ('Networking', 'Traffic flow', 'Routes reachable', '', ''),
    ('Storage', 'Volume lifecycle', 'PVCs provision and bind', '', ''),
    ('Security', 'Login and RBAC', 'Auth and roles enforced', '', ''),
    ('Applications', 'App deployment', 'Workloads run end-to-end', '', ''),
    ('Resilience', 'Node failure', 'Workloads recover', '', ''),
    ('Scaling', 'Add capacity', 'New node joins cluster', '', ''),
    ('Backup', 'Protect and restore', 'Data intact after restore', '', ''),
    ('Monitoring', 'Alert pipeline', 'Alerts delivered', '', ''),
    ('Migration', 'VM migration', 'VMs operational', '', ''),
    ('Upgrade', 'Version bump', 'Upgrade succeeds cleanly', '', ''),
])

# Phase 10
add_phase_heading('Phase 10: Closeout')

add_section_heading('Deliverables')
doc.add_paragraph(
    'The POC concludes with two formal deliverables:'
)
p1 = doc.add_paragraph()
run = p1.add_run('1. Completed checklist')
run.bold = True
run.font.name = FONT_TEXT
p1.add_run(
    ' \u2014 this document, filled in with status and notes for every item tested.'
).font.name = FONT_TEXT

p2 = doc.add_paragraph()
run = p2.add_run('2. Closeout meeting')
run.bold = True
run.font.name = FONT_TEXT
p2.add_run(
    ' \u2014 a readout of all POC findings, outcomes, gaps, and '
    'recommendations presented to stakeholders.'
).font.name = FONT_TEXT
doc.add_paragraph()

add_section_heading('Findings summary')
add_checklist_table([
    'Successful tests documented',
    'Failures documented with root cause and resolution',
    'Platform gaps identified (features not yet available)',
    'Infrastructure gaps identified (hardware, network, storage)',
    'Application gaps identified (app-specific constraints)',
    'POC-only limitations noted (not relevant to production)',
    'Lessons learned recorded for production planning',
    'Final go/no-go recommendation delivered and agreed',
])

add_section_heading('Operational readiness of customer team')
doc.add_paragraph('The customer team has demonstrated the ability to:')
add_checklist_table([
    'Perform routine cluster administration',
    'Diagnose and resolve common issues',
    'Execute cluster upgrades',
    'Add or remove cluster capacity',
    'Run backup and restore procedures',
    'Deploy new applications to the platform',
])

add_section_heading('Formal approval')
add_signoff_table()

# Follow-Up
add_phase_heading('Follow-up: Sizing and proposal')
doc.add_paragraph(
    'Upon successful completion of the POC, the next step is a sizing and proposal '
    'process that translates POC findings into a production-ready architecture and '
    'commercial agreement.'
)
add_checklist_table([
    'Compute sizing documented (CPU, memory, disk per node role)',
    'Storage architecture and capacity plan defined',
    'Network topology and segmentation documented',
    'High-availability and DR strategy outlined',
    'Backup retention and RPO/RTO targets set',
    'Security hardening steps identified',
    'Alerting and on-call strategy documented',
    'Operational ownership model agreed (who runs what)',
    'Hardware bill of materials finalized',
    'Network allocation documented (subnets, IPs, firewall rules)',
    'Red Hat entitlements and subscription counts confirmed',
    'External dependencies cataloged (storage, load balancers, DNS)',
    'Commercial proposal delivered to customer',
])

# Save
output_dir = os.path.join(os.path.dirname(__file__), '..', 'docs', 'assets', 'downloads')
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'poc-checklist.docx')
doc.save(output_path)
print(f'Generated {output_path}')
